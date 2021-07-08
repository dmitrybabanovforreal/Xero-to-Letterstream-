#! /usr/bin/python3

from flask import Flask, request, Response, redirect
import requests, config, datetime, docx, os, traceback
from docx.shared import Pt
from base64 import b64encode
from docx2pdf import convert


app = Flask(__name__)


@app.route(config.authSlug)
def authorization():

    try:
        token = open('data/token').read()
    except:
        token, refreshToken = '', ''

    # if the is no token, launch the process of the app authorization
    if token == '':
        # check if user returned from the authorization page with the code
        # https://developer.xero.com/documentation/oauth2/auth-flow#redirect
        authCode = request.args.get('code')
        if authCode:
            # send the code to get the token
            # https://developer.xero.com/documentation/oauth2/auth-flow#code
            s = config.id + ":" + config.secret
            headers = {'Authorization': 'Basic ' + str(b64encode(s.encode("utf-8")), "utf-8")}
            data = {
                'grant_type': 'authorization_code',
                'code': authCode,
                'redirect_uri': config.appAddress + config.authSlug
            }
            response = requests.post('https://identity.xero.com/connect/token', headers=headers, data=data)
            token = response.json()['access_token']
            refreshToken = response.json()['refresh_token']

            file = open('data/token', 'w')
            file.write(token)
            file.close()

            file = open('data/refreshToken', 'w')
            file.write(refreshToken)
            file.close()

            # get the tennant ID
            # https://developer.xero.com/documentation/oauth2/auth-flow#connections
            headers = {'Authorization': 'Bearer ' + token}
            response = requests.get('https://api.xero.com/connections', headers=headers)
            tenantId = response.json()[0]['tenantId']
            file = open('data/tenantId', 'w')
            file.write(tenantId)
            file.close()

            return f'The app is authorized. You can now call {config.appAddress}{config.invoicesSlug} to trigger the ' \
                   f'invoices processing.\n\nYou can pass startdate and enddate parameters in YYYY-MM-DD format to ' \
                   f'specify the invoices dates.\n\nBy default it processes all invoices generated yesterday'

        else:
            # todo specify the correct scope in the config file
            # send the authorization request to get the code
            # https://developer.xero.com/documentation/oauth2/auth-flow#authorize
            redirectUrl = f'https://login.xero.com/identity/connect/authorize?response_type=code&client_id={config.id}&redirect_uri={config.appAddress}{config.authSlug}&scope={config.scope}'
            return redirect(redirectUrl)


@app.route(config.invoicesSlug)
def process_invoices():
    # check that all the credentials are available
    try:
        token = open('data/token').read()
        refreshToken = open('data/refreshToken').read()
        tenantId = open('data/tenantId').read()
    except:
        return f'Apparently, the app is not authorized yet. Open {config.appAddress}{config.authSlug} to start the ' \
               f'authorization process'

    # since Xero token only lasts for 30 minutes and we plan to use their API daily,
    # we need almost always start with refreshing the token
    # https://developer.xero.com/documentation/oauth2/auth-flow#refresh
    s = config.id + ":" + config.secret
    headers = {'Authorization': 'Basic ' + str(b64encode(s.encode("utf-8")), "utf-8")}
    data = {'grant_type': 'refresh_token', 'refresh_token': refreshToken}
    response = requests.post('https://identity.xero.com/connect/token', headers=headers, data=data)

    # try refreshing the token and if really expired, save the new refresh token to use later
    try:
        token = response.json()['access_token']
        refreshToken = response.json()['refresh_token']

        file = open('data/access_token', 'w')
        file.write(token)
        file.close()

        file = open('data/refreshToken', 'w')
        file.write(refreshToken)
        file.close()
    except:
        pass

    # get the optional start/end date parameters
    startdate = request.args.get('startdate', default=(datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%Y-%m-%d"))
    enddate = request.args.get('enddate', default=(datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%Y-%m-%d"))

    if startdate is None or enddate is None:
        startdate = (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%Y-%m-%d")  # yesterday
        enddate = (datetime.datetime.now() - datetime.timedelta(days=1)).strftime("%Y-%m-%d")  # yesterday

    # date range string to pass later to Xero API
    dateRange = f"Date >= DateTime({','.join(startdate.split('-'))}) AND Date <= DateTime({','.join(enddate.split('-'))})"

    # call the Invoices API to get the data
    # using paging to enable getting the line items for each invoice
    # https://developer.xero.com/documentation/api/invoices
    i = 0
    invoices = []
    while True:
        i += 1
        headers = {'Authorization': 'Bearer ' + token, 'xero-tenant-id': tenantId, 'Accept': 'application/json'}
        params = {'where': dateRange, 'page': i}
        response = requests.get('https://api.xero.com/api.xro/2.0/Invoices/', headers=headers, params=params)
        try:
            if len(response.json()['Invoices']) == 0:
                break
        except:
            break
        invoices += response.json()['Invoices']

    # todo remove saving the invoices data to file
    file = open('invoices json', 'w')
    file.write(str(invoices))
    file.close()

    os.mkdir('data/invoices')

    for invoice in invoices:
        headers = {'Authorization': 'Bearer ' + token, 'xero-tenant-id': tenantId, 'Accept': 'application/pdf'}
        response = requests.get('https://api.xero.com/api.xro/2.0/Invoices/' + invoice['InvoiceID'], headers=headers, params=params)
        file = open(os.path.join('data/invoices', invoice['InvoiceID'] + '.pdf'), 'wb')
        file.write(response.content)
        file.close()

    for invoice in invoices:
        # get the details of the customer contact for each invoice; remove the demo contact
        response = requests.get('https://api.xero.com/api.xro/2.0/Contacts/' + invoice['Contact']['ContactID'])
        contact = response.json()['Contacts'][0]

        addressText = '\n' + '\n'.join([
            contact['Name'],
            '\n'.join([
                contact['Addresses'][0].get('AddressLine1', ''),
                contact['Addresses'][0].get('AddressLine2', ''),
                contact['Addresses'][0].get('AddressLine3', ''),
                contact['Addresses'][0].get('AddressLine4', '')
            ]).strip('\n'),
            contact['Addresses'][0].get('City', '') + ' ' + contact['Addresses'][0].get('Region', ''),
            contact['Addresses'][0].get('PostalCode', '')
        ])

        # todo add the fork: summary invoice or the complete one

        # open and fill in the docx file of the "total only" invoice template
        doc = docx.Document('templates/Invoice total template.docx')
        doc.tables[0].rows[3].cells[1].paragraphs[0].runs[0].text = datetime.datetime.strptime(invoice['DateString'], '%Y-%m-%dT%H:%M:%S').strftime('%b %d %Y')
        doc.tables[0].rows[3].cells[2].paragraphs[0].runs[0].text = datetime.datetime.strptime(invoice['DueDateString'], '%Y-%m-%dT%H:%M:%S').strftime('%b %d %Y')
        doc.tables[0].rows[3].cells[0].text = addressText
        doc.tables[0].rows[3].cells[0].paragraphs[0].runs[0].font.size = Pt(9)

        doc.tables[0].rows[5].cells[1].paragraphs[0].runs[0].text = invoice['InvoiceNumber']
        doc.tables[0].rows[5].cells[2].paragraphs[0].runs[0].text = invoice.get('Reference', '')

        doc.tables[1].rows[0].cells[2].paragraphs[0].runs[0].text = invoice['SubTotal']
        # todo add Total State Sales Tax, can't find this field in the demo invoice
        # todo add Total County Sales Tax, can't find this field in the demo invoice
        doc.tables[1].rows[3].cells[2].paragraphs[0].runs[0].text = invoice['Total']

        doc.tables[2].rows[0].cells[1].paragraphs[0].runs[0].text = contact['Name']
        doc.tables[2].rows[1].cells[1].paragraphs[0].runs[0].text = invoice['InvoiceNumber']
        doc.tables[2].rows[3].cells[1].paragraphs[0].runs[0].text = invoice['AmountDue']
        doc.tables[2].rows[4].cells[1].paragraphs[0].runs[0].text = datetime.datetime.strptime(invoice['DueDateString'], '%Y-%m-%dT%H:%M:%S').strftime('%d %b %Y')

        doc.save(invoice['InvoiceNumber'] + ' summary.docx')
        convert(invoice['InvoiceNumber'] + ' summary.docx', invoice['InvoiceNumber'] + ' summary.pdf')
        os.remove(invoice['InvoiceNumber'] + ' summary.docx')

        # todo add sending the file to the API


        # open and fill in the docx file of the "complete" invoice template
        doc = docx.Document('templates/Invoice complete template.docx')
        doc.tables[0].rows[0].cells[1].paragraphs[1].runs[0].text = datetime.datetime.strptime(invoice['DateString'],'%Y-%m-%dT%H:%M:%S').strftime('%b %d %Y')
        doc.tables[0].rows[0].cells[1].paragraphs[3].runs[0].text = invoice['InvoiceNumber']
        doc.tables[0].rows[0].cells[1].paragraphs[5].runs[0].text = invoice.get('Reference', '')
        doc.tables[0].rows[0].cells[0].paragraphs[1].runs[0].text = addressText

        for item in invoice['LineItems']:
            row = doc.tables[1].add_row()
            row.cells[0].text = item['Description']
            row.cells[1].text = item['Quantity']
            row.cells[2].text = item['UnitAmount']
            row.cells[3].text = item['LineAmount']
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                cell.paragraphs[0].runs[0].font.name = 'Arial'

        doc.tables[2].rows[0].cells[2].paragraphs[0].runs[0].text = invoice['SubTotal']
        # todo add Total State Sales Tax, can't find this field in the demo invoice
        # todo add Total County Sales Tax, can't find this field in the demo invoice
        doc.tables[2].rows[2].cells[2].paragraphs[0].runs[0].text = invoice['Total']
        doc.tables[2].rows[3].cells[2].paragraphs[0].runs[0].text = invoice['AmountCredited']
        doc.tables[2].rows[4].cells[2].paragraphs[0].runs[0].text = invoice['AmountDue']

        doc.paragraphs[9].runs[0].text = 'Due Date: ' + datetime.datetime.strptime(invoice['DueDateString'], '%Y-%m-%dT%H:%M:%S').strftime('%b %d %Y')

        doc.tables[3].rows[2].cells[3].paragraphs[0].runs[0].text = contact['Name']
        doc.tables[3].rows[3].cells[3].paragraphs[0].runs[0].text = invoice['InvoiceNumber']
        doc.tables[3].rows[4].cells[3].paragraphs[0].runs[0].text = invoice['AmountDue']
        doc.tables[3].rows[5].cells[3].paragraphs[0].runs[0].text = datetime.datetime.strptime(invoice['DueDateString'], '%Y-%m-%dT%H:%M:%S').strftime('%b %d %Y')

        doc.save(invoice['InvoiceNumber'] + ' detailed.docx')
        convert(invoice['InvoiceNumber'] + ' detailed.docx', invoice['InvoiceNumber'] + ' detailed.pdf')
        os.remove(invoice['InvoiceNumber'] + ' detailed.docx')

        # todo add sending the file to the API

    return Response(status=200)


@app.route('/')
def respond():
    return 'Hello world'